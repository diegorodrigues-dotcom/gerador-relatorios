import streamlit as st
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io

import reportlab
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.pdfgen import canvas

# Configuração da página Web
st.set_page_config(page_title="Gerador de Relatórios - Kärcher", layout="wide", page_icon="⚙️")

st.title("⚙️ Gerador de Relatórios Técnicos - Padrão Kärcher")
st.subheader("Lavadoras de Alta Pressão e Equipamentos Motorizados")

st.markdown("---")

def set_cell_background(cell, fill_hex):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def add_field(paragraph, field_type):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_type
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r = run._r
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)

def safe_float(val):
    """Converte texto digitado para float de forma segura para cálculo da média"""
    try:
        if not val:
            return 0.0
        return float(str(val).replace(',', '.'))
    except ValueError:
        return 0.0

# Classe Canvas para numerar páginas no PDF dinamicamente
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super(NumberedCanvas, self).__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.setFont("Helvetica", 8)
        self.setFillColor(colors.HexColor("#333333"))
        
        st_text = f"ST {self.codigo_st}" if self.codigo_st else "ST"
        page_text = f"Folha {self._pageNumber} de {page_count}"
        item_text = self.item_testado if self.item_testado else ""
        
        self.drawString(36, 25, st_text)
        self.drawCentredString(306, 25, page_text)
        self.drawRightString(576, 25, item_text)

# 1. CABEÇALHO / IDENTIFICAÇÃO GERAL
st.header("1. Informações Gerais do Ensaio")
c1, c2, c3 = st.columns(3)

with c1:
    codigo_st = st.text_input("ST", value="", placeholder="Ex: 001941")
    objetivo = st.text_area("Objetivo do Teste", value="", height=80, placeholder="Digite o objetivo do teste...")

with c2:
    tecnico = st.text_input("Responsável", value="", placeholder="Digite seu nome...")
    normas = st.text_area("Critério de Aprovação / Normas", value="", height=80, placeholder="Digite os critérios / normas...")

with c3:
    item_testado = st.text_input("Item Testado", value="", placeholder="O que você está testando?...")
    data_ensaio = st.text_input("Data do Teste", value="", placeholder="Ex: 01/01/2026")

conclusao_texto = st.text_area("Conclusão / Parecer Técnico Geral", 
    value="", 
    height=120,
    placeholder="Digite aqui a conclusão e parecer técnico geral...")

st.markdown("---")

# 2. CADASTRO DINÂMICO DE AMOSTRAS E MEDIÇÕES
st.header("2. Cadastro de Amostras (Medições, Fotos e Defeitos)")

c_quant, c_rpm_opt = st.columns([1, 1])

with c_quant:
    num_amostras = st.number_input("Quantidade de Amostras", min_value=1, value=1)

with c_rpm_opt:
    incluir_rpm = st.selectbox("Equipamento possui medição de RPM?", ["Não", "Sim"]) == "Sim"

amostras_dados = []

for idx in range(int(num_amostras)):
    # GARANTE EXIBIÇÃO CORRETA: "Amostra 1", "Amostra 2", etc.
    st.markdown(f"## Amostra {idx+1}")
    
    col_a, col_b, col_c = st.columns(3)
    with col_a:
        sample_id = st.text_input(f"Sample ID", value="", placeholder="Ex: AM 1", key=f"id_{idx}")
        voltagem_conexao = st.text_input("Tensão", value="", placeholder="127V / 220V", key=f"volt_{idx}")
    with col_b:
        col_p1, col_p2 = st.columns([1, 1])
        with col_p1:
            partiu_frio = st.selectbox("Partiu a frio?", ["Sim", "Não"], key=f"p_status_{idx}")
        with col_p2:
            tensao_partida = st.text_input("Tensão de Partida", value="", placeholder="Ex: 94V", key=f"p_tensao_{idx}")
        
        partida = f"{partiu_frio} ({tensao_partida})" if tensao_partida else partiu_frio
        horas_ensaio = st.text_input("Tempo de Teste / Horas", value="", placeholder="Ex: 116 h", key=f"h_{idx}")
    with col_c:
        defeitos_texto = st.text_area("Lista de Falhas", value="", placeholder="Digite as falhas encontradas...", key=f"def_{idx}", height=80)

    st.write(f"**Parâmetros de Teste Funcional - {sample_id if sample_id else f'Amostra {idx+1}'}:**")
    
    if incluir_rpm:
        m1, m2, m3, m4, m5, m6 = st.columns(6)
    else:
        m1, m2, m3, m4, m5 = st.columns(5)
    
    with m1:
        st.caption("Tensão (V)")
        v30 = st.text_input("30s (V)", value="", placeholder="Ex: 126.2", key=f"v30_{idx}")
        v3m = st.text_input("3min (V)", value="", placeholder="Ex: 128.0", key=f"v3m_{idx}")
        v5m = st.text_input("5min (V)", value="", placeholder="Ex: 127.1", key=f"v5m_{idx}")

    with m2:
        st.caption("Potência (kW)")
        p30 = st.text_input("30s (kW)", value="", placeholder="Ex: 1.53", key=f"p30_{idx}")
        p3m = st.text_input("3min (kW)", value="", placeholder="Ex: 1.55", key=f"p3m_{idx}")
        p5m = st.text_input("5min (kW)", value="", placeholder="Ex: 1.50", key=f"p5m_{idx}")

    with m3:
        st.caption("Pressão bico")
        pr30 = st.text_input("30s (Bar)", value="", placeholder="Ex: 91.9", key=f"pr30_{idx}")
        pr3m = st.text_input("3min (Bar)", value="", placeholder="Ex: 94.0", key=f"pr3m_{idx}")
        pr5m = st.text_input("5min (Bar)", value="", placeholder="Ex: 94.5", key=f"pr5m_{idx}")

    with m4:
        st.caption("Vazão (l/h)")
        vz30 = st.text_input("30s (l/h)", value="", placeholder="Ex: 293", key=f"vz30_{idx}")
        vz3m = st.text_input("3min (l/h)", value="", placeholder="Ex: 296", key=f"vz3m_{idx}")
        vz5m = st.text_input("5min (l/h)", value="", placeholder="Ex: 297", key=f"vz5m_{idx}")

    with m5:
        st.caption("Corrente (A)")
        i30 = st.text_input("30s (A)", value="", placeholder="Ex: 12.68", key=f"i30_{idx}")
        i3m = st.text_input("3min (A)", value="", placeholder="Ex: 12.60", key=f"i3m_{idx}")
        i5m = st.text_input("5min (A)", value="", placeholder="Ex: 12.38", key=f"i5m_{idx}")

    rpm30, rpm3m, rpm5m, mrpm = "", "", "", ""
    if incluir_rpm:
        with m6:
            st.caption("RPM (rpm)")
            rpm30 = st.text_input("30s (RPM)", value="", placeholder="Ex: 3450", key=f"rpm30_{idx}")
            rpm3m = st.text_input("3min (RPM)", value="", placeholder="Ex: 3420", key=f"rpm3m_{idx}")
            rpm5m = st.text_input("5min (RPM)", value="", placeholder="Ex: 3410", key=f"rpm5m_{idx}")

    # CÁLCULO DE MÉDIAS
    num_v = [safe_float(v30), safe_float(v3m), safe_float(v5m)]
    num_p = [safe_float(p30), safe_float(p3m), safe_float(p5m)]
    num_pr = [safe_float(pr30), safe_float(pr3m), safe_float(pr5m)]
    num_vz = [safe_float(vz30), safe_float(vz3m), safe_float(vz5m)]
    num_i = [safe_float(i30), safe_float(i3m), safe_float(i5m)]

    mv = round(sum(num_v)/3, 1) if any(num_v) else ""
    mp = round(sum(num_p)/3, 2) if any(num_p) else ""
    mpr = round(sum(num_pr)/3, 1) if any(num_pr) else ""
    mvz = round(sum(num_vz)/3, 0) if any(num_vz) else ""
    mi = round(sum(num_i)/3, 2) if any(num_i) else ""

    if incluir_rpm:
        num_rpm = [safe_float(rpm30), safe_float(rpm3m), safe_float(rpm5m)]
        calc_mrpm = round(sum(num_rpm)/3, 0) if any(num_rpm) else ""
        mrpm = str(int(calc_mrpm)) if calc_mrpm != "" else ""

    fotos_uploaded = st.file_uploader(f"Anexar Imagens para {sample_id if sample_id else f'Amostra {idx+1}'}", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key=f"foto_{idx}")

    amostras_dados.append({
        "sample_id": sample_id,
        "voltagem_conexao": voltagem_conexao,
        "partida": partida,
        "horas": horas_ensaio,
        "defeitos": defeitos_texto,
        "fotos": fotos_uploaded,
        "v30": v30, "v3m": v3m, "v5m": v5m, "mv": str(mv) if mv != "" else "",
        "p30": p30, "p3m": p3m, "p5m": p5m, "mp": str(mp) if mp != "" else "",
        "pr30": pr30, "pr3m": pr3m, "pr5m": pr5m, "mpr": str(mpr) if mpr != "" else "",
        "vz30": vz30, "vz3m": vz3m, "vz5m": vz5m, "mvz": str(int(mvz)) if mvz != "" else "",
        "i30": i30, "i3m": i3m, "i5m": i5m, "mi": str(mi) if mi != "" else "",
        "rpm30": rpm30, "rpm3m": rpm3m, "rpm5m": rpm5m, "mrpm": mrpm
    })
    st.markdown("---")

# ----------------------------------------------------
# SEÇÃO DE BOTÕES DE GERAÇÃO (DOCX e PDF)
# ----------------------------------------------------
st.header("3. Opções de Download do Relatório")
col_btn1, col_btn2 = st.columns(2)

# OPÇÃO 1: GERAR EM WORD (.DOCX)
with col_btn1:
    if st.button("🚀 GERAR RELATÓRIO WORD (.DOCX)", type="primary", use_container_width=True):
        doc = docx.Document()

        for section in doc.sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            
            footer = section.footer
            p_ft = footer.paragraphs[0]
            p_ft.text = ""
            
            tbl_ft = footer.add_table(1, 3, Inches(6.9))
            tbl_ft.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            c_st_ft, c_page_ft, c_item_ft = tbl_ft.rows[0].cells[0], tbl_ft.rows[0].cells[1], tbl_ft.rows[0].cells[2]
            c_st_ft.width, c_page_ft.width, c_item_ft.width = Inches(2.3), Inches(2.3), Inches(2.3)
            
            p_st_ft = c_st_ft.paragraphs[0]
            r_st = p_st_ft.add_run(f"ST {codigo_st}" if codigo_st else "ST")
            r_st.font.size = Pt(8.5)
            r_st.font.color.rgb = RGBColor(50, 50, 50)
            
            p_page = c_page_ft.paragraphs[0]
            p_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_p1 = p_page.add_run("Folha ")
            r_p1.font.size = Pt(8.5)
            r_p1.font.color.rgb = RGBColor(50, 50, 50)
            add_field(p_page, 'PAGE')
            r_p2 = p_page.add_run(" de ")
            r_p2.font.size = Pt(8.5)
            r_p2.font.color.rgb = RGBColor(50, 50, 50)
            add_field(p_page, 'NUMPAGES')
            
            p_item = c_item_ft.paragraphs[0]
            p_item.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r_it = p_item.add_run(item_testado if item_testado else "")
            r_it.font.size = Pt(8.5)
            r_it.font.color.rgb = RGBColor(50, 50, 50)

        tbl_hdr = doc.add_table(rows=1, cols=2)
        tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_left, c_right = tbl_hdr.rows[0].cells[0], tbl_hdr.rows[0].cells[1]
        c_left.width, c_right.width = Inches(4.5), Inches(2.0)

        p_dept = c_left.paragraphs[0]
        r_dept = p_dept.add_run("Departamento de testes e desenvolvimentos")
        r_dept.bold = True
        r_dept.font.size = Pt(11)

        p_logo = c_right.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_logo = p_logo.add_run("KÄRCHER")
        r_logo.bold = True
        r_logo.font.size = Pt(18)

        p_main_title = doc.add_paragraph()
        p_main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_main_title.paragraph_format.space_before = Pt(18)
        p_main_title.paragraph_format.space_after = Pt(24)
        r_title = p_main_title.add_run("Relatório de teste")
        r_title.bold = True
        r_title.font.size = Pt(26)

        tbl_top = doc.add_table(rows=1, cols=2)
        tbl_top.style = 'Table Grid'
        tbl_top.alignment = WD_TABLE_ALIGNMENT.CENTER
        c_st_label, c_st_val = tbl_top.rows[0].cells[0], tbl_top.rows[0].cells[1]
        c_st_label.width, c_st_val.width = Inches(1.2), Inches(5.3)
        
        p0 = c_st_label.paragraphs[0]
        r0 = p0.add_run("ST")
        r0.bold, r0.font.size = True, Pt(14)

        p1 = c_st_val.paragraphs[0]
        r1 = p1.add_run(codigo_st)
        r1.bold, r1.font.size = True, Pt(14)

        doc.add_paragraph()

        tbl_meta = doc.add_table(rows=6, cols=2)
        tbl_meta.style = 'Table Grid'
        tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER

        meta_info = [
            ("Teste realizado por", tecnico),
            ("Data", data_ensaio),
            ("Item testado", item_testado),
            ("Quantidade", str(num_amostras)),
            ("Objetivo do teste", objetivo),
            ("Critério de aprovação", normas)
        ]

        for i, (k, v) in enumerate(meta_info):
            row = tbl_meta.rows[i]
            c0, c1 = row.cells[0], row.cells[1]
            c0.width, c1.width = Inches(2.2), Inches(4.3)
            set_cell_background(c0, "F2F2F2")
            c0.paragraphs[0].add_run(k).bold = True
            c1.paragraphs[0].add_run(v)

        doc.add_paragraph()

        doc.add_paragraph().add_run("Conclusão").bold = True
        doc.add_paragraph(conclusao_texto)
        doc.add_paragraph()

        doc.add_paragraph().add_run("1- Teste funcional de Parâmetros").bold = True

        for am in amostras_dados:
            p_sub = doc.add_paragraph()
            p_sub.add_run(f"Máquina com conexão {am['voltagem_conexao']}").bold = True
            
            num_cols = 9 if incluir_rpm else 8
            tbl = doc.add_table(rows=5, cols=num_cols)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

            if incluir_rpm:
                headers = ["Tempo de teste:", f"Partida a frio {am['partida']}", "Tensão (V) - 60 Hz", "Potência absorvida (kW)", "Pressão com bico", "Vazão (l/h)", "Corrente (A)", "RPM (rpm)", "Sample ID"]
                rows_data = [
                    ("30s", "OK", str(am["v30"]), str(am["p30"]), str(am["pr30"]), str(am["vz30"]), str(am["i30"]), str(am["rpm30"]), am["sample_id"]),
                    ("3 min", "", str(am["v3m"]), str(am["p3m"]), str(am["pr3m"]), str(am["vz3m"]), str(am["i3m"]), str(am["rpm3m"]), ""),
                    ("5 min", "", str(am["v5m"]), str(am["p5m"]), str(am["pr5m"]), str(am["vz5m"]), str(am["i5m"]), str(am["rpm5m"]), ""),
                    ("Média", "", str(am["mv"]), str(am["mp"]), str(am["mpr"]), str(am["mvz"]), str(am["mi"]), str(am["mrpm"]), "")
                ]
            else:
                headers = ["Tempo de teste:", f"Partida a frio {am['partida']}", "Tensão (V) - 60 Hz", "Potência absorvida (kW)", "Pressão com bico", "Vazão (l/h)", "Corrente (A)", "Sample ID"]
                rows_data = [
                    ("30s", "OK", str(am["v30"]), str(am["p30"]), str(am["pr30"]), str(am["vz30"]), str(am["i30"]), am["sample_id"]),
                    ("3 min", "", str(am["v3m"]), str(am["p3m"]), str(am["pr3m"]), str(am["vz3m"]), str(am["i3m"]), ""),
                    ("5 min", "", str(am["v5m"]), str(am["p5m"]), str(am["pr5m"]), str(am["vz5m"]), str(am["i5m"]), ""),
                    ("Média", "", str(am["mv"]), str(am["mp"]), str(am["mpr"]), str(am["mvz"]), str(am["mi"]), "")
                ]
            
            hdr_row = tbl.rows[0]
            for col_i, h_text in enumerate(headers):
                cell = hdr_row.cells[col_i]
                set_cell_background(cell, "A6A6A6")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(h_text)
                r.bold = True
                r.font.size = Pt(8.0 if incluir_rpm else 8.5)
                r.font.color.rgb = RGBColor(0, 0, 0)

            for r_idx, r_vals in enumerate(rows_data):
                row = tbl.rows[r_idx + 1]
                for c_idx, val in enumerate(r_vals):
                    cell = row.cells[c_idx]
                    
                    if r_vals[0] == "Média" and c_idx in ([0] + list(range(2, num_cols-1))):
                        set_cell_background(cell, "A6A6A6")
                        
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(val)
                    r.font.size = Pt(8.5 if incluir_rpm else 9.0)
                    if r_vals[0] == "Média":
                        r.bold = True

            doc.add_paragraph()

        doc.add_paragraph().add_run("2- Durabilidade conforme norma KN 082.023 cap. 4.7.1").bold = True

        for am in amostras_dados:
            doc.add_paragraph().add_run(f"• {am['sample_id']} ({am['voltagem_conexao']})").bold = True
            
            if am["fotos"]:
                cols_count = min(len(am["fotos"]), 3)
                tbl_ft = doc.add_table(rows=1, cols=cols_count)
                tbl_ft.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                for f_i, f_file in enumerate(am["fotos"]):
                    if f_i < 3:
                        cell = tbl_ft.rows[0].cells[f_i]
                        p_img = cell.paragraphs[0]
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_stream = io.BytesIO(f_file.read())
                        p_img.add_run().add_picture(img_stream, width=Inches(1.8))

            doc.add_paragraph(f"Após {am['horas']} foram identificadas as falhas / defeitos:")
            doc.add_paragraph(am["defeitos"])
            doc.add_paragraph()

        doc.add_paragraph().add_run("Resumo das informações de defeitos e durabilidade apresentadas pelas amostras").bold = True
        
        tbl_res = doc.add_table(rows=len(amostras_dados)+1, cols=4)
        tbl_res.style = 'Table Grid'
        
        headers_res = ["Amostra", "Tempo de teste", "Vida útil esperada", "Falhas apresentadas"]
        for c_i, h_txt in enumerate(headers_res):
            cell = tbl_res.rows[0].cells[c_i]
            set_cell_background(cell, "A6A6A6")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(h_txt).bold = True

        for r_i, am in enumerate(amostras_dados):
            row = tbl_res.rows[r_i+1]
            row.cells[0].paragraphs[0].add_run(am["sample_id"])
            row.cells[1].paragraphs[0].add_run(am["horas"])
            row.cells[2].paragraphs[0].add_run("60 h")
            row.cells[3].paragraphs[0].add_run("Identificadas no ensaio")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.success("✅ Relatório Word (.docx) gerado com sucesso!")
        st.download_button(
            label="📥 Baixar Documento Word (.docx)",
            data=buffer,
            file_name=f"ST {codigo_st if codigo_st else '000'} - Karcher {item_testado if item_testado else 'Relatorio'}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# OPÇÃO 2: GERAR EM PDF (.PDF)
with col_btn2:
    if st.button("📄 GERAR RELATÓRIO PDF (.PDF)", type="secondary", use_container_width=True):
        pdf_buffer = io.BytesIO()
        pdf_doc = SimpleDocTemplate(pdf_buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=40)
        
        styles = getSampleStyleSheet()
        style_title = ParagraphStyle(name='TitleStyle', fontName='Helvetica-Bold', fontSize=12, leading=14)
        style_main_title = ParagraphStyle(name='MainTitleStyle', fontName='Helvetica-Bold', fontSize=22, leading=26, alignment=1)
        style_normal = ParagraphStyle(name='NormalStyle', fontName='Helvetica', fontSize=9, leading=11)
        style_bold = ParagraphStyle(name='BoldStyle', fontName='Helvetica-Bold', fontSize=9, leading=11)
        
        elements = []

        hdr_data = [[
            Paragraph("<b>Departamento de testes e desenvolvimentos</b>", style_bold),
            Paragraph("<b>KÄRCHER</b>", ParagraphStyle(name='RLogo', fontName='Helvetica-Bold', fontSize=16, alignment=2))
        ]]
        t_hdr = Table(hdr_data, colWidths=[380, 160])
        elements.append(t_hdr)
        elements.append(Spacer(1, 15))

        elements.append(Paragraph("Relatório de teste", style_main_title))
        elements.append(Spacer(1, 15))

        data_st = [[Paragraph("<b>ST</b>", style_title), Paragraph(codigo_st, style_title)]]
        t_st = Table(data_st, colWidths=[60, 480])
        t_st.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 1, colors.black),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        elements.append(t_st)
        elements.append(Spacer(1, 10))

        meta_data = [
            [Paragraph("<b>Teste realizado por</b>", style_bold), Paragraph(tecnico, style_normal)],
            [Paragraph("<b>Data</b>", style_bold), Paragraph(data_ensaio, style_normal)],
            [Paragraph("<b>Item testado</b>", style_bold), Paragraph(item_testado, style_normal)],
            [Paragraph("<b>Quantidade</b>", style_bold), Paragraph(str(num_amostras), style_normal)],
            [Paragraph("<b>Objetivo do teste</b>", style_bold), Paragraph(objetivo, style_normal)],
            [Paragraph("<b>Critério de aprovação</b>", style_bold), Paragraph(normas, style_normal)],
        ]
        t_meta = Table(meta_data, colWidths=[150, 390])
        t_meta.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (0,0), (0,-1), colors.HexColor("#F2F2F2")),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ]))
        elements.append(t_meta)
        elements.append(Spacer(1, 10))

        elements.append(Paragraph("<b>Conclusão</b>", style_bold))
        elements.append(Paragraph(conclusao_texto, style_normal))
        elements.append(Spacer(1, 10))

        elements.append(Paragraph("<b>1- Teste funcional de Parâmetros</b>", style_bold))
        elements.append(Spacer(1, 5))

        for am in amostras_dados:
            elements.append(Paragraph(f"<b>Máquina com conexão {am['voltagem_conexao']}</b>", style_normal))
            
            if incluir_rpm:
                headers_pdf = ["Tempo:", f"Partida {am['partida']}", "Tensão (V)", "Potência (kW)", "Pressão", "Vazão (l/h)", "Corrente (A)", "RPM", "Sample ID"]
                param_data = [headers_pdf]
                param_data.append(["30s", "OK", str(am["v30"]), str(am["p30"]), str(am["pr30"]), str(am["vz30"]), str(am["i30"]), str(am["rpm30"]), am["sample_id"]])
                param_data.append(["3 min", "", str(am["v3m"]), str(am["p3m"]), str(am["pr3m"]), str(am["vz3m"]), str(am["i3m"]), str(am["rpm3m"]), ""])
                param_data.append(["5 min", "", str(am["v5m"]), str(am["p5m"]), str(am["pr5m"]), str(am["vz5m"]), str(am["i5m"]), str(am["rpm5m"]), ""])
                param_data.append(["Média", "", str(am["mv"]), str(am["mp"]), str(am["mpr"]), str(am["mvz"]), str(am["mi"]), str(am["mrpm"]), ""])
                col_widths = [50, 65, 55, 60, 55, 55, 60, 55, 85]
            else:
                headers_pdf = ["Tempo:", f"Partida {am['partida']}", "Tensão (V)", "Potência (kW)", "Pressão", "Vazão (l/h)", "Corrente (A)", "Sample ID"]
                param_data = [headers_pdf]
                param_data.append(["30s", "OK", str(am["v30"]), str(am["p30"]), str(am["pr30"]), str(am["vz30"]), str(am["i30"]), am["sample_id"]])
                param_data.append(["3 min", "", str(am["v3m"]), str(am["p3m"]), str(am["pr3m"]), str(am["vz3m"]), str(am["i3m"]), ""])
                param_data.append(["5 min", "", str(am["v5m"]), str(am["p5m"]), str(am["pr5m"]), str(am["vz5m"]), str(am["i5m"]), ""])
                param_data.append(["Média", "", str(am["mv"]), str(am["mp"]), str(am["mpr"]), str(am["mvz"]), str(am["mi"]), ""])
                col_widths = [55, 75, 65, 65, 65, 65, 65, 85]

            t_param = Table(param_data, colWidths=col_widths)
            t_param.setStyle(TableStyle([
                ('GRID', (0,0), (-1,-1), 0.5, colors.black),
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#A6A6A6")),
                ('BACKGROUND', (0,-1), (-2,-1), colors.HexColor("#A6A6A6")),
                ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
                ('FONTSIZE', (0,0), (-1,-1), 7.5 if incluir_rpm else 8),
            ]))
            elements.append(t_param)
            elements.append(Spacer(1, 8))

        elements.append(Paragraph("<b>2- Durabilidade conforme norma KN 082.023 cap. 4.7.1</b>", style_bold))
        for am in amostras_dados:
            elements.append(Paragraph(f"• {am['sample_id']} ({am['voltagem_conexao']})", style_normal))
            
            if am["fotos"]:
                img_list = []
                for f_file in am["fotos"][:3]:
                    img_stream = io.BytesIO(f_file.read())
                    f_file.seek(0)
                    rl_img = RLImage(img_stream, width=110, height=110)
                    img_list.append(rl_img)
                if img_list:
                    t_imgs = Table([img_list])
                    elements.append(t_imgs)
                    elements.append(Spacer(1, 5))

            elements.append(Paragraph(f"Após {am['horas']} foram identificadas as falhas / defeitos:", style_normal))
            elements.append(Paragraph(am['defeitos'], style_normal))
            elements.append(Spacer(1, 8))

        elements.append(Paragraph("<b>Resumo das informações de defeitos e durabilidade apresentadas pelas amostras</b>", style_bold))
        elements.append(Spacer(1, 5))
        
        res_data = [["Amostra", "Tempo de teste", "Vida útil esperada", "Falhas apresentadas"]]
        for am in amostras_dados:
            res_data.append([am["sample_id"], am["horas"], "60 h", "Identificadas no ensaio"])
            
        t_res = Table(res_data, colWidths=[100, 100, 100, 240])
        t_res.setStyle(TableStyle([
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#A6A6A6")),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('FONTSIZE', (0,0), (-1,-1), 8),
        ]))
        elements.append(t_res)

        def make_canvas(*args, **kwargs):
            c = NumberedCanvas(*args, **kwargs)
            c.codigo_st = codigo_st
            c.item_testado = item_testado
            return c

        pdf_doc.build(elements, canvasmaker=make_canvas)
        pdf_buffer.seek(0)

        st.success("✅ Relatório PDF (.pdf) gerado com sucesso!")
        st.download_button(
            label="📥 Baixar Documento PDF (.pdf)",
            data=pdf_buffer,
            file_name=f"ST {codigo_st if codigo_st else '000'} - Karcher {item_testado if item_testado else 'Relatorio'}.pdf",
            mime="application/pdf",
            use_container_width=True
        )
